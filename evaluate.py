"""
Evaluation and reporting tool for combined LLM classification results.

This script:
- Optionally filters highly similar spans with conflicting gold labels.
- Computes precision, recall, and F1-score per label for each model.
- Prints results to console, or exports them as a formatted Word (.docx) report.
- Copies a tab-delimited version of the report to the clipboard for easy pasting.

Typical usage:
--------------
python evaluate.py \
    --input combined_results.json \
    --output metrics.json \
    --filter --threshold 0.75 \
    --filtered_output filtered_results \
    --use_docx
"""

import json
import argparse
from difflib import SequenceMatcher
from docx import Document
import pyperclip

from sklearn.metrics import classification_report
from promt_templates import FULL_LABEL_NAMES


def filter_results(results, threshold):
    """
    Remove conflicting spans within a policy.

    For each policy_id, if two spans are at least `threshold` similar
    (by difflib.SequenceMatcher) but have different gold labels,
    both spans are dropped.

    Parameters
    ----------
    results : list of dict
        Records with keys including 'policy_id', 'span', 'gold'.
    threshold : float
        Similarity ratio (0–1) above which spans are considered duplicates.

    Returns
    -------
    list of dict
        Filtered results.
    """
    to_drop = set()
    policy_groups = {}
    for idx, r in enumerate(results):
        pid = r.get('policy_id')
        policy_groups.setdefault(pid, []).append(idx)

    for pid, indices in policy_groups.items():
        for i in range(len(indices)):
            for j in range(i + 1, len(indices)):
                idx_i = indices[i]
                idx_j = indices[j]
                r_i = results[idx_i]
                r_j = results[idx_j]
                span_i = str(r_i['span'])
                span_j = str(r_j['span'])
                gold_i = r_i['gold']
                gold_j = r_j['gold']
                if gold_i != gold_j:
                    ratio = SequenceMatcher(None, span_i, span_j).ratio()
                    if ratio >= threshold:
                        to_drop.update([idx_i, idx_j])
    return [r for idx, r in enumerate(results) if idx not in to_drop]


def evaluate(combined, output_path=None):
    """
    Compute classification metrics for each model.

    Parameters
    ----------
    combined : dict
        Mapping model_id → list of records with 'gold' and 'pred'.
    output_path : str or None
        Path to save metrics as JSON, if provided.

    Returns
    -------
    dict
        Metrics dictionary keyed by model_id.
    """
    all_metrics = {}
    for model_id, results in combined.items():
        y_true = [r["gold"] for r in results]
        y_pred = [r["pred"] for r in results]
        report = classification_report(
            y_true,
            y_pred,
            labels=FULL_LABEL_NAMES,
            zero_division=0,
            output_dict=True
        )
        all_metrics[model_id] = report
    if output_path:
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(all_metrics, f, indent=2, ensure_ascii=False)
    return all_metrics


def compute_distribution(results):
    """
    Compute dataset size and distribution of gold labels.

    Parameters
    ----------
    results : list of dict
        Records with 'gold' label.

    Returns
    -------
    total : int
        Total number of records.
    dist : dict
        Mapping label → count.
    """
    total = len(results)
    dist = {label: 0 for label in FULL_LABEL_NAMES}
    for r in results:
        label = r.get('gold')
        if label in dist:
            dist[label] += 1
    return total, dist


def print_report(metrics: dict, total: int, dist: dict):
    """
    Print label distribution and classification metrics to console.
    """
    print(f"Total sections: {total}")
    print("Label distribution (gold):")
    print(f"{'Label':<50s}{'Count':>8s}")
    for label, count in dist.items():
        print(f"{label:<50s}{count:8d}")
    print("\n")

    for model_id, rpt in metrics.items():
        print(f"=== Model: {model_id} ===")
        print(f"{'Label':<50s}{'P':>8s}{'R':>8s}{'F1':>8s}")
        for label in FULL_LABEL_NAMES:
            scores = rpt.get(label, {})
            p = scores.get("precision", 0) * 100
            r = scores.get("recall", 0) * 100
            f = scores.get("f1-score", 0) * 100
            print(f"{label:<50s}{p:8.0f}{r:8.0f}{f:8.0f}")
        mac = rpt.get("macro avg", {})
        p, r, f = mac.get("precision", 0) * 100, mac.get("recall", 0) * 100, mac.get("f1-score", 0) * 100
        print(f"{'Macro average':<50s}{p:8.0f}{r:8.0f}{f:8.0f}")
        print("-" * 74)


def write_report(metrics: dict, total: int, dist: dict, doc_path: str = "report.docx"):
    """
    Write the classification report to a Word file and copy text to clipboard.

    Parameters
    ----------
    metrics : dict
        Metrics dictionary from evaluate().
    total : int
        Number of records.
    dist : dict
        Gold label distribution.
    doc_path : str
        Destination .docx file path.
    """
    doc = Document()
    lines = []

    doc.add_heading("Classification Report", level=1)
    lines.append("Classification Report\n")

    doc.add_paragraph(f"Total sections: {total}")
    lines.append(f"Total sections:\t{total}\n")

    doc.add_paragraph("Label distribution (gold):")
    lines.append("Label distribution (gold):")

    dist_table = doc.add_table(rows=1, cols=2)
    dist_table.style = 'Table Grid'
    dist_table.autofit = True
    hdr_cells = dist_table.rows[0].cells
    hdr_cells[0].text = "Label"
    hdr_cells[1].text = "Count"
    lines.append("Label\tCount")
    for label, count in dist.items():
        row_cells = dist_table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = str(count)
        lines.append(f"{label}\t{count}")
    lines.append("")

    for model_id, rpt in metrics.items():
        doc.add_heading(f"Model: {model_id}", level=2)
        lines.append(f"Model:\t{model_id}")
        lines.append("Label\tP\tR\tF1")

        tbl = doc.add_table(rows=1 + len(FULL_LABEL_NAMES) + 1, cols=4)
        tbl.style = 'Table Grid'
        tbl.autofit = True

        hdr = tbl.rows[0].cells
        hdr[0].text = "Label"
        hdr[1].text = "P"
        hdr[2].text = "R"
        hdr[3].text = "F1"

        for i, label in enumerate(FULL_LABEL_NAMES, start=1):
            scores = rpt.get(label, {})
            p = scores.get("precision", 0) * 100
            r = scores.get("recall", 0) * 100
            f = scores.get("f1-score", 0) * 100
            cells = tbl.rows[i].cells
            cells[0].text = label
            cells[1].text = f"{p:.0f}"
            cells[2].text = f"{r:.0f}"
            cells[3].text = f"{f:.0f}"
            lines.append(f"{label}\t{p:.0f}\t{r:.0f}\t{f:.0f}")

        mac = rpt.get("macro avg", {})
        p = mac.get("precision", 0) * 100
        r = mac.get("recall", 0) * 100
        f = mac.get("f1-score", 0) * 100
        mac_row = tbl.rows[len(FULL_LABEL_NAMES) + 1].cells
        mac_row[0].text = "Macro average"
        mac_row[1].text = f"{p:.0f}"
        mac_row[2].text = f"{r:.0f}"
        mac_row[3].text = f"{f:.0f}"
        lines.append(f"Macro average\t{p:.0f}\t{r:.0f}\t{f:.0f}")
        lines.append("")

    doc.save(doc_path)
    pyperclip.copy("\n".join(lines))


if __name__ == '__main__':
    parser = argparse.ArgumentParser(
        description='Evaluate combined LLM results with optional span filtering'
    )
    parser.add_argument('--input', required=True, help='Path to combined_results.json')
    parser.add_argument('--output', default=None, help='Optional: path to save metrics JSON')
    parser.add_argument('--threshold', type=float, default=0.75, help='Similarity threshold (0–1) for filtering')
    parser.add_argument('--filter', action='store_true', help='Enable filtering of spans based on threshold')
    parser.add_argument('--filtered_output', default=None, help='Optional: path to save filtered results JSON')
    parser.add_argument('--use_docx', action='store_true', help='Write report to .docx instead of console')
    args = parser.parse_args()

    with open(args.input, 'r', encoding='utf-8') as f:
        combined = json.load(f)

    if args.filter:
        print(f"Filtering spans with threshold {args.threshold}...")
        for model_id, results in combined.items():
            combined[model_id] = filter_results(results, args.threshold)
        if args.filtered_output:
            path = args.filtered_output + "/" + args.input.split('/')[-1]
            with open(path, 'w', encoding='utf-8') as f:
                json.dump(combined, f, indent=2, ensure_ascii=False)
            print(f"Filtered results saved to {args.filtered_output}")
    else:
        print("Skipping span filtering...")

    first_model = next(iter(combined))
    total, dist = compute_distribution(combined[first_model])
    metrics = evaluate(combined, args.output)

    if args.use_docx:
        write_report(metrics, total, dist)
    else:
        print_report(metrics, total, dist)
